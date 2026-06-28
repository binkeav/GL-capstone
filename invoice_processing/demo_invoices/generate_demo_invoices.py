from __future__ import annotations

import json
from datetime import date, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT_DIR = Path(__file__).resolve().parent


def today_minus(days: int) -> str:
    return (date.today() - timedelta(days=days)).isoformat()


PASS_INVOICE = {
    "invoice_number": "INV-DEMO-001",
    "invoice_date": today_minus(10),
    "due_date": today_minus(-20),
    "po_number": "PO-1001",
    "payment_terms": "Net 30",
    "vendor_name": "Acme Supplies India",
    "vendor_tax_id": "29ABCDE1234F1Z5",
    "customer_name": "Demo Buyer India",
    "customer_tax_id": "29ZZZZZ9999Z1Z5",
    "subtotal": 1200,
    "tax": 216,
    "shipping": 0,
    "discounts": 0,
    "total": 1416,
    "currency": "INR",
    "order_items": [
        {
            "line_no": 1,
            "description": "Laptop Stand",
            "qty": 1,
            "unit": "pcs",
            "unit_price": 1000,
            "net_amount": 1000,
            "tax_rate": 18,
            "gross_amount": 1180,
        },
        {
            "line_no": 2,
            "description": "USB Cable",
            "qty": 2,
            "unit": "pcs",
            "unit_price": 100,
            "net_amount": 200,
            "tax_rate": 18,
            "gross_amount": 236,
        },
    ],
}

REVIEW_INVOICE = {
    **PASS_INVOICE,
    "invoice_number": "INV-DEMO-002",
    "po_number": "PO-404",
    "total": 1400,
}

CRITICAL_INVOICE = {
    **PASS_INVOICE,
    "invoice_number": "INV-DEMO-003",
    "invoice_date": today_minus(20),
    "due_date": today_minus(-10),
    "vendor_name": "Blocked Vendor India",
    "vendor_tax_id": "29BBBBB1234B1Z5",
}


SAMPLES = [
    ("01_straight_through_acme", "Straight-through processing demo", PASS_INVOICE),
    ("02_human_review_mismatch", "Human review demo", REVIEW_INVOICE),
    ("03_critical_blocked_vendor", "Critical rejection demo", CRITICAL_INVOICE),
]


def write_json(name: str, payload: dict[str, Any]) -> None:
    (OUT_DIR / f"{name}.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")


def write_docx(name: str, title: str, payload: dict[str, Any]) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Invoice Number: {payload['invoice_number']}")
    doc.add_paragraph(f"Invoice Date: {payload['invoice_date']}")
    doc.add_paragraph(f"Due Date: {payload['due_date']}")
    doc.add_paragraph(f"PO Number: {payload['po_number']}")
    doc.add_paragraph(f"Payment Terms: {payload['payment_terms']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Vendor Name: {payload['vendor_name']}")
    doc.add_paragraph(f"Vendor GSTIN: {payload['vendor_tax_id']}")
    doc.add_paragraph(f"Customer Name: {payload['customer_name']}")
    doc.add_paragraph(f"Customer GSTIN: {payload['customer_tax_id']}")

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["Line", "Description", "Qty", "Unit", "Unit Price", "Net", "Tax Rate", "Gross"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for item in payload["order_items"]:
        cells = table.add_row().cells
        values = [
            item["line_no"],
            item["description"],
            item["qty"],
            item["unit"],
            item["unit_price"],
            item["net_amount"],
            item["tax_rate"],
            item["gross_amount"],
        ]
        for idx, value in enumerate(values):
            cells[idx].text = str(value)

    doc.add_paragraph("")
    doc.add_paragraph(f"Subtotal: {payload['subtotal']}")
    doc.add_paragraph(f"Tax: {payload['tax']}")
    doc.add_paragraph(f"Shipping: {payload['shipping']}")
    doc.add_paragraph(f"Discounts: {payload['discounts']}")
    doc.add_paragraph(f"Total: {payload['total']}")
    doc.add_paragraph(f"Currency: {payload['currency']}")

    doc.save(OUT_DIR / f"{name}.docx")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for name, title, payload in SAMPLES:
        write_json(name, payload)
        write_docx(name, title, payload)
    index = {
        "samples": [
            {
                "name": name,
                "docx": f"{name}.docx",
                "json": f"{name}.json",
                "expected_demo_state": expected,
            }
            for (name, _, _), expected in zip(
                SAMPLES,
                ["STRAIGHT_THROUGH_PROCESSING", "HUMAN_IN_THE_LOOP", "CRITICAL_REJECTION"],
            )
        ]
    }
    (OUT_DIR / "index.json").write_text(json.dumps(index, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
