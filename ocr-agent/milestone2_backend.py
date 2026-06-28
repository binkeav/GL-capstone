import json
import logging
import os
import atexit
import ast
import subprocess
import tempfile
import threading
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import cv2
import easyocr
import httpx
import numpy as np
import pypdfium2 as pdfium
from dotenv import load_dotenv
from openai import OpenAI

#using local models for easyocr (craft_mlt_25k.pth and english_g2.pth), and local qwen 8b model due to huggingface blockage on network
#qwen model downloaded from https://www.kaggle.com/datasets/nareshmeena12/qwen3-vl-8b-instruct/data

logger = logging.getLogger(__name__)


def _configure_ocr_logging(workspace_root: Path) -> Path:
    log_file = os.getenv("INVOICE_PROCESSING_LOG_FILE")
    path = Path(log_file) if log_file else workspace_root / "outputs" / "ocr_agent.log"
    if not path.is_absolute():
        path = workspace_root / path
    path.parent.mkdir(parents=True, exist_ok=True)
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)
    if not any(
        isinstance(handler, logging.FileHandler)
        and Path(getattr(handler, "baseFilename", "")) == path
        for handler in root_logger.handlers
    ):
        handler = logging.FileHandler(path, encoding="utf-8")
        handler.setFormatter(
            logging.Formatter(
                "%(asctime)s %(levelname)s %(name)s [%(process)d:%(threadName)s] %(message)s"
            )
        )
        root_logger.addHandler(handler)
    return path

class Milestone1NotebookAPI:
    """Backend API adapted from 03_milestone1_easyocr_only notebook logic."""

    SCALAR_FIELDS = [
        "invoice_number",
        "invoice_date",
        "due_date",
        "po_number",
        "payment_terms",
        "vendor_name",
        "vendor_tax_id",
        "customer_name",
        "customer_tax_id",
        "subtotal",
        "tax",
        "shipping",
        "discounts",
        "total",
        "currency",
    ]
    ORDER_ITEM_FIELDS = [
        "line_no",
        "description",
        "qty",
        "unit",
        "unit_price",
        "net_amount",
        "tax_rate",
        "gross_amount",
    ]

    def __init__(self, workspace_root: Path):
        self.workspace_root = workspace_root
        self.dataset_root = workspace_root / "Datasets"
        self.output_dir = workspace_root / "outputs"
        self.output_dir.mkdir(exist_ok=True)
        self.log_file = _configure_ocr_logging(workspace_root)
        logger.info("Initializing OCR backend workspace_root=%s log_file=%s", workspace_root, self.log_file)

        self.easyocr_model_dir = workspace_root
        self.easyocr_user_network_dir = workspace_root / "user_network"
        self.craft_model_path = self.easyocr_model_dir / "craft_mlt_25k.pth"
        self.english_model_path = self.easyocr_model_dir / "english_g2.pth"

        if not self.craft_model_path.exists() or not self.english_model_path.exists():
            raise FileNotFoundError(
                "Required EasyOCR model files not found: craft_mlt_25k.pth and english_g2.pth"
            )

        self.reader = easyocr.Reader(
            ["en"],
            gpu=False,
            model_storage_directory=str(self.easyocr_model_dir),
            user_network_directory=str(self.easyocr_user_network_dir),
            detect_network="craft",
            recog_network="english_g2",
            download_enabled=False,
            verbose=False,
        )

        env_file = workspace_root / ".env"
        load_dotenv(env_file)
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.openai_api_base = os.getenv(
            "OPENAI_API_BASE", "https://aibe.mygreatlearning.com/openai/v1"
        )
        self.model_name = "gpt-4o-mini"
        self.field_extractor_mode = os.getenv("FIELD_EXTRACTOR_MODE", "auto").strip().lower()
        if self.field_extractor_mode not in {"auto", "gpt", "qwen"}:
            self.field_extractor_mode = "auto"

        self.qwen_model_dir = self._resolve_qwen_model_dir(workspace_root)
        self.qwen_model_name = self._qwen_display_name(self.qwen_model_dir)
        self._qwen_model = None
        self._qwen_processor = None
        self._qwen_device = "cpu"
        self._qwen_load_error: Optional[str] = None
        self.qwen_device_preference = os.getenv("QWEN_DEVICE", "auto").strip().lower()
        self.qwen_max_new_tokens = self._env_int("QWEN_MAX_NEW_TOKENS", 512, 32, 1500)
        self.qwen_max_time_seconds = self._env_float("QWEN_MAX_TIME_SECONDS", 90.0, 5.0, 600.0)
        self.qwen_retry_max_time_seconds = self._env_float("QWEN_RETRY_MAX_TIME_SECONDS", 240.0, 10.0, 900.0)
        self.qwen_max_ocr_chars = self._env_int("QWEN_MAX_OCR_CHARS", 6000, 1000, 20000)
        logger.info(
            "OCR extraction config mode=%s openai_configured=%s qwen_model=%s qwen_dir=%s qwen_device_preference=%s max_new_tokens=%s max_time=%s retry_max_time=%s max_ocr_chars=%s",
            self.field_extractor_mode,
            bool(self.openai_api_key),
            self.qwen_model_name,
            self.qwen_model_dir,
            self.qwen_device_preference,
            self.qwen_max_new_tokens,
            self.qwen_max_time_seconds,
            self.qwen_retry_max_time_seconds,
            self.qwen_max_ocr_chars,
        )

        self.layout_worker_python = workspace_root / ".venv-layout" / "bin" / "python"
        self.layout_worker_script = workspace_root / "scripts" / "layout_worker.py"
        self._layout_proc = None
        self._layout_lock = threading.Lock()
        self._layout_error: Optional[str] = None

        self.openai_client = None
        if self.openai_api_key:
            http_client = httpx.Client(verify=False)
            self.openai_client = OpenAI(
                api_key=self.openai_api_key,
                base_url=self.openai_api_base,
                http_client=http_client,
            )

    @staticmethod
    def _env_int(name: str, default: int, min_value: int, max_value: int) -> int:
        try:
            value = int(os.getenv(name, str(default)))
        except ValueError:
            return default
        return max(min_value, min(max_value, value))

    @staticmethod
    def _env_float(name: str, default: float, min_value: float, max_value: float) -> float:
        try:
            value = float(os.getenv(name, str(default)))
        except ValueError:
            return default
        return max(min_value, min(max_value, value))

    def _ensure_layout_worker(self) -> bool:
        if self._layout_proc is not None and self._layout_proc.poll() is None:
            return True
        if self._layout_error is not None:
            return False
        if not self.layout_worker_python.exists() or not self.layout_worker_script.exists():
            self._layout_error = (
                "layout_worker_missing: "
                f"python={self.layout_worker_python.exists()} "
                f"worker={self.layout_worker_script.exists()}"
            )
            return False
        try:
            self._layout_proc = subprocess.Popen(
                [
                    str(self.layout_worker_python),
                    str(self.layout_worker_script),
                    "--batch",
                    "--score-thresh",
                    "0.3",
                    "--device",
                    "cpu",
                ],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.DEVNULL,
                text=True,
                bufsize=1,
            )
            atexit.register(self._close_layout_worker)
            return True
        except Exception as exc:
            self._layout_error = f"layout_worker_start_failed: {exc}"
            return False

    def __del__(self):
        self._close_layout_worker()

    def _close_layout_worker(self) -> None:
        try:
            if self._layout_proc is not None and self._layout_proc.poll() is None:
                self._layout_proc.stdin.close()
                self._layout_proc.terminate()
                self._layout_proc.wait(timeout=5)
        except Exception:
            pass

    def detect_layout_regions(self, image_rgb: np.ndarray) -> Dict[str, Any]:
        if not self._ensure_layout_worker():
            return {
                "status": "layout_worker_unavailable",
                "regions": [],
                "note": self._layout_error,
            }

        fd, tmp_path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        try:
            bgr = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2BGR)
            cv2.imwrite(tmp_path, bgr)
            with self._layout_lock:
                self._layout_proc.stdin.write(tmp_path + "\n")
                self._layout_proc.stdin.flush()
                line = self._layout_proc.stdout.readline()
            if not line:
                return {"status": "layout_worker_dead", "regions": []}
            return json.loads(line)
        except Exception as exc:
            return {"status": "layout_worker_error", "regions": [], "note": str(exc)}
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    @staticmethod
    def _sort_regions_reading_order(regions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        def _key(r: Dict[str, Any]) -> Tuple[int, int]:
            x1, y1, _, _ = r.get("bbox", [0, 0, 0, 0])
            return int(y1), int(x1)

        return sorted(regions, key=_key)

    def ocr_image_layout_aware(self, image_rgb: np.ndarray) -> Tuple[str, float, int, Dict[str, Any]]:
        # Baseline page OCR (milestone-1 behavior).
        full_text, full_conf, full_det, full_pre_meta = self.easyocr_with_preprocessing_variants(
            image_rgb,
            exhaustive=True,
        )

        layout = self.detect_layout_regions(image_rgb)
        regions = layout.get("regions", []) if isinstance(layout, dict) else []
        if not regions:
            return full_text, full_conf, full_det, {
                "layout_status": layout.get("status", "unknown") if isinstance(layout, dict) else "unknown",
                "layout_regions": 0,
                "layout_crop_chunks": 0,
                "preprocessing": full_pre_meta,
            }

        h, w = image_rgb.shape[:2]
        crop_texts: List[str] = []
        crop_confs: List[float] = []
        crop_dets = 0

        ordered = self._sort_regions_reading_order(regions)
        # Keep runtime bounded for interactive uploads.
        for idx, r in enumerate(ordered[:30], start=1):
            bbox = r.get("bbox", [])
            if len(bbox) != 4:
                continue
            x1, y1, x2, y2 = [int(v) for v in bbox]
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(w, x2), min(h, y2)
            if x2 - x1 < 12 or y2 - y1 < 12:
                continue

            crop = image_rgb[y1:y2, x1:x2]
            crop_text, crop_conf, crop_det, _ = self.easyocr_with_preprocessing_variants(crop)
            if crop_det <= 0 or not crop_text.strip():
                continue

            crop_texts.append(f"[REGION {idx}]\n{crop_text}")
            crop_confs.append(crop_conf)
            crop_dets += crop_det

        if not crop_texts:
            return full_text, full_conf, full_det, {
                "layout_status": layout.get("status", "unknown"),
                "layout_regions": len(regions),
                "layout_crop_chunks": 0,
                "preprocessing": full_pre_meta,
            }

        merged_text = (
            "=== FULL PAGE OCR ===\n"
            + full_text
            + "\n\n=== LAYOUT REGION OCR ===\n"
            + "\n\n".join(crop_texts)
        )
        conf_values = [v for v in [full_conf] + crop_confs if isinstance(v, (int, float))]
        merged_conf = float(sum(conf_values) / len(conf_values)) if conf_values else full_conf
        merged_det = int(full_det + crop_dets)

        return merged_text, merged_conf, merged_det, {
            "layout_status": layout.get("status", "unknown"),
            "layout_regions": len(regions),
            "layout_crop_chunks": len(crop_texts),
            "preprocessing": full_pre_meta,
        }

    def _empty_fields(self) -> Dict[str, Any]:
        data = {k: None for k in self.SCALAR_FIELDS}
        data["order_items"] = []
        return data

    @staticmethod
    def _get_nested(data: Dict[str, Any], path: str) -> Any:
        current: Any = data
        for part in path.split("."):
            if not isinstance(current, dict):
                return None
            current = current.get(part)
        return current

    def _first_value(self, data: Dict[str, Any], keys: List[str]) -> Any:
        for key in keys:
            value = self._get_nested(data, key) if "." in key else data.get(key)
            if value not in (None, ""):
                return value
        return None

    def _normalize_extracted_fields(self, parsed: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        out = self._empty_fields()
        if not isinstance(parsed, dict):
            return out

        scalar_aliases = {
            "invoice_number": ["invoice_number", "invoice_no", "invoice_id", "invoice"],
            "invoice_date": ["invoice_date", "date_of_issue", "issue_date", "date"],
            "due_date": ["due_date", "payment_due_date"],
            "po_number": ["po_number", "po_no", "purchase_order", "purchase_order_number"],
            "payment_terms": ["payment_terms", "terms"],
            "vendor_name": ["vendor_name", "supplier_name", "seller_name", "vendor.name", "supplier.name", "seller.name"],
            "vendor_tax_id": [
                "vendor_tax_id",
                "vendor_gstin",
                "supplier_tax_id",
                "supplier_gstin",
                "seller_tax_id",
                "seller_gstin",
                "vendor.tax_id",
                "vendor.gstin",
                "supplier.tax_id",
                "supplier.gstin",
                "seller.tax_id",
                "seller.gstin",
            ],
            "customer_name": ["customer_name", "buyer_name", "client_name", "customer.name", "buyer.name", "client.name"],
            "customer_tax_id": [
                "customer_tax_id",
                "customer_gstin",
                "buyer_tax_id",
                "buyer_gstin",
                "client_tax_id",
                "client_gstin",
                "customer.tax_id",
                "customer.gstin",
                "buyer.tax_id",
                "buyer.gstin",
                "client.tax_id",
                "client.gstin",
            ],
            "subtotal": ["subtotal", "sub_total", "net_total", "summary.subtotal", "summary.net_worth", "summary.net_amount"],
            "tax": ["tax", "tax_total", "total_tax", "vat_amount", "summary.tax", "summary.vat_amount", "summary.tax_total"],
            "shipping": ["shipping", "shipping_amount", "freight", "summary.shipping"],
            "discounts": ["discounts", "discount", "discount_amount", "summary.discounts", "summary.discount"],
            "total": ["total", "grand_total", "amount_due", "gross_total", "summary.total", "summary.gross_worth", "summary.grand_total"],
            "currency": ["currency", "currency_code"],
        }

        for key in self.SCALAR_FIELDS:
            out[key] = self._first_value(parsed, scalar_aliases.get(key, [key]))

        raw_items = self._first_value(parsed, ["order_items", "line_items", "items", "invoice_items"])
        items: List[Dict[str, Any]] = []
        if isinstance(raw_items, list):
            for item in raw_items:
                if not isinstance(item, dict):
                    continue
                item_aliases = {
                    "line_no": ["line_no", "line_number", "item_no", "item_number"],
                    "description": ["description", "item_description", "name", "product", "service"],
                    "qty": ["qty", "quantity", "amount"],
                    "unit": ["unit", "unit_measure", "uom"],
                    "unit_price": ["unit_price", "net_price", "price", "rate"],
                    "net_amount": ["net_amount", "net_worth", "line_total", "amount"],
                    "tax_rate": ["tax_rate", "vat_pct", "tax_percent", "gst_rate"],
                    "gross_amount": ["gross_amount", "gross_worth", "total", "line_gross_total"],
                }
                items.append({k: self._first_value(item, item_aliases.get(k, [k])) for k in self.ORDER_ITEM_FIELDS})
        out["order_items"] = items
        return out

    def quality_metrics(self, image_rgb: np.ndarray) -> Dict[str, float]:
        gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
        blur = float(cv2.Laplacian(gray, cv2.CV_64F).var())
        brightness = float(np.mean(gray))
        contrast = float(np.std(gray))
        return {"blur_var": blur, "brightness": brightness, "contrast": contrast}

    def preprocess_for_ocr(self, image_rgb: np.ndarray, *, invert: bool = False) -> np.ndarray:
        gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
        if invert:
            gray = cv2.bitwise_not(gray)
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
        th = cv2.adaptiveThreshold(
            denoised,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            5,
        )
        return cv2.cvtColor(th, cv2.COLOR_GRAY2RGB)

    def easyocr_with_preprocessing_variants(
        self,
        image_rgb: np.ndarray,
        *,
        exhaustive: bool = False,
    ) -> Tuple[str, float, int, Dict[str, Any]]:
        dark_background = self._looks_like_dark_background(image_rgb)
        variants = self._ocr_preprocessing_variants(image_rgb, exhaustive=exhaustive or dark_background)

        variant_results = []
        for name, image in variants:
            text, conf, det = self.easyocr_on_image_array(image)
            score = (det * 2.0) + (conf * 3.0) + min(len(text), 2000) / 1000.0
            variant_results.append(
                {
                    "name": name,
                    "text": text,
                    "confidence": conf,
                    "detections": det,
                    "score": score,
                }
            )

        ranked = sorted(variant_results, key=lambda item: item["score"], reverse=True)
        selected = [item for item in ranked if item["detections"] > 0 and item["text"].strip()][:3]
        if not selected and ranked:
            selected = [ranked[0]]

        merged_text = self._merge_ocr_variant_texts(selected)
        merged_conf_values = [item["confidence"] for item in selected if isinstance(item["confidence"], (int, float))]
        merged_conf = float(sum(merged_conf_values) / len(merged_conf_values)) if merged_conf_values else 0.0
        merged_det = int(sum(item["detections"] for item in selected))
        scores = {
            item["name"]: {
                "detections": item["detections"],
                "confidence": item["confidence"],
                "score": round(item["score"], 4),
            }
            for item in variant_results
        }

        return merged_text, merged_conf, merged_det, {
            "selected_variant": selected[0]["name"] if selected else "none",
            "merged_variants": [item["name"] for item in selected],
            "variants": scores,
            "dark_background_detected": dark_background,
        }

    def _ocr_preprocessing_variants(
        self,
        image_rgb: np.ndarray,
        *,
        exhaustive: bool = False,
    ) -> List[Tuple[str, np.ndarray]]:
        variants = [
            ("adaptive", self.preprocess_for_ocr(image_rgb, invert=False)),
            ("adaptive_inverted", self.preprocess_for_ocr(image_rgb, invert=True)),
        ]
        if exhaustive:
            variants.extend(
                [
                    ("clahe", self._clahe_for_ocr(image_rgb, invert=False)),
                    ("clahe_inverted", self._clahe_for_ocr(image_rgb, invert=True)),
                    ("gray", self._gray_for_ocr(image_rgb, invert=False)),
                    ("gray_inverted", self._gray_for_ocr(image_rgb, invert=True)),
                ]
            )
        return variants

    def _clahe_for_ocr(self, image_rgb: np.ndarray, *, invert: bool = False) -> np.ndarray:
        gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
        if invert:
            gray = cv2.bitwise_not(gray)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        enhanced = cv2.bilateralFilter(enhanced, 5, 50, 50)
        return cv2.cvtColor(enhanced, cv2.COLOR_GRAY2RGB)

    @staticmethod
    def _gray_for_ocr(image_rgb: np.ndarray, *, invert: bool = False) -> np.ndarray:
        gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
        if invert:
            gray = cv2.bitwise_not(gray)
        return cv2.cvtColor(gray, cv2.COLOR_GRAY2RGB)

    @staticmethod
    def _merge_ocr_variant_texts(variant_results: List[Dict[str, Any]]) -> str:
        seen = set()
        chunks = []
        for result in variant_results:
            lines = []
            for line in str(result.get("text") or "").splitlines():
                clean = " ".join(line.strip().split())
                if not clean:
                    continue
                key = clean.lower()
                if key in seen:
                    continue
                seen.add(key)
                lines.append(line.strip())
            if lines:
                chunks.append(f"=== OCR VARIANT: {result['name']} ===\n" + "\n".join(lines))
        return "\n\n".join(chunks)

    @staticmethod
    def _looks_like_dark_background(image_rgb: np.ndarray) -> bool:
        gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
        brightness = float(np.mean(gray))
        dark_ratio = float(np.mean(gray < 90))
        bright_ratio = float(np.mean(gray > 180))
        return brightness < 120 and dark_ratio > 0.45 and bright_ratio > 0.02

    def easyocr_on_image_array(self, image_rgb: np.ndarray) -> Tuple[str, float, int]:
        results = self.reader.readtext(image_rgb)
        texts, confs = [], []
        for item in results:
            if len(item) >= 3:
                txt = str(item[1]).strip()
                if txt:
                    texts.append(txt)
                    confs.append(float(item[2]))
        text = "\n".join(texts)
        avg_conf = float(sum(confs) / len(confs)) if confs else 0.0
        return text, avg_conf, len(texts)

    def _extraction_error_fields(self, reason: str, detail: Optional[str] = None) -> Dict[str, Any]:
        logger.error("Field extraction error reason=%s detail=%s", reason, detail)
        out = self._empty_fields()
        out["extraction_error"] = reason
        if detail:
            out["extraction_error_detail"] = detail
        return out

    @staticmethod
    def _balanced_json_candidates(content: str) -> List[str]:
        candidates = []
        for start, char in enumerate(content):
            if char not in "{[":
                continue
            stack = [char]
            in_string = False
            escape = False
            for idx in range(start + 1, len(content)):
                current = content[idx]
                if escape:
                    escape = False
                    continue
                if current == "\\":
                    escape = True
                    continue
                if current == '"':
                    in_string = not in_string
                    continue
                if in_string:
                    continue
                if current in "{[":
                    stack.append(current)
                elif current in "}]":
                    if not stack:
                        break
                    opening = stack.pop()
                    if (opening, current) not in {("{", "}"), ("[", "]")}:
                        break
                    if not stack:
                        candidates.append(content[start : idx + 1])
                        break
        return candidates

    @staticmethod
    def _parse_json_candidate(candidate: str) -> Any:
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            try:
                return ast.literal_eval(candidate)
            except (ValueError, SyntaxError):
                return None

    def _parse_model_json(self, content: str) -> Optional[Dict[str, Any]]:
        stripped = content.strip()
        candidates = [stripped]

        if "```" in stripped:
            parts = stripped.split("```")
            for idx, part in enumerate(parts):
                block = part.strip()
                if idx % 2 == 1:
                    if block.lower().startswith("json"):
                        block = block[4:].strip()
                    candidates.append(block)

        candidates.extend(self._balanced_json_candidates(stripped))

        for candidate in candidates:
            parsed = self._parse_json_candidate(candidate.strip())
            if isinstance(parsed, dict):
                return parsed
            if isinstance(parsed, list):
                for item in parsed:
                    if isinstance(item, dict):
                        return item
        return None

    def _parse_json_or_error(self, content: str, reason: str) -> Dict[str, Any]:
        parsed = self._parse_model_json(content)
        if parsed is not None:
            return self._normalize_extracted_fields(parsed)
        return self._extraction_error_fields(reason, content[:500])

    def _looks_like_partial_json(self, content: str) -> bool:
        text = (content or "").strip()
        if len(text) < 20 and any(char in text for char in "{["):
            return True
        return text.count("{") > text.count("}") or text.count("[") > text.count("]")

    def _qwen_ocr_context(self, ocr_text: str) -> str:
        text = ocr_text.strip()
        if len(text) <= self.qwen_max_ocr_chars:
            return text

        head_chars = int(self.qwen_max_ocr_chars * 0.7)
        tail_chars = self.qwen_max_ocr_chars - head_chars
        return (
            text[:head_chars]
            + "\n\n[... middle OCR text omitted for local Qwen speed ...]\n\n"
            + text[-tail_chars:]
        )

    def _resolve_qwen_model_dir(self, workspace_root: Path) -> Path:
        configured_dir = os.getenv("QWEN_MODEL_DIR", "").strip()
        if configured_dir:
            model_dir = Path(configured_dir).expanduser()
            if not model_dir.is_absolute():
                model_dir = workspace_root / model_dir
            return model_dir

        for candidate in (
            "qwen3-vl-8b",
            "Qwen3-VL-8B-Instruct",
            "qwen3-vl-8b-instruct",
            "qwen-vl",
        ):
            model_dir = workspace_root / candidate
            if model_dir.exists():
                return model_dir

        return workspace_root / "qwen3-vl-8b"

    def _qwen_display_name(self, model_dir: Path) -> str:
        config_path = model_dir / "config.json"
        try:
            with config_path.open("r", encoding="utf-8") as config_file:
                config = json.load(config_file)
            architecture = (config.get("architectures") or [None])[0]
            model_type = config.get("model_type")
            if architecture:
                return f"{architecture} (local)"
            if model_type:
                return f"{model_type} (local)"
        except Exception:
            pass
        return f"{model_dir.name} (local)"

    def _qwen_device_candidates(self) -> List[str]:
        preference = self.qwen_device_preference
        if preference in {"cpu", "cuda", "mps"}:
            return [preference]
        try:
            import torch

            candidates = []
            if torch.cuda.is_available():
                candidates.append("cuda")
            if hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
                candidates.append("mps")
            candidates.append("cpu")
            return candidates
        except Exception:
            return ["cpu"]

    @staticmethod
    def is_policy_block(message: str) -> bool:
        markers = [
            "zscaler",
            "violates compliance category",
            "posting content to this website is not allowed",
            "<!doctype html",
            "dlp policy",
            "internet security by zscaler",
        ]
        low = message.lower()
        return any(m in low for m in markers)

    def extract_fields_gpt4omini(self, ocr_text: str) -> Dict[str, Any]:
        if self.openai_client is None:
            return self._extraction_error_fields("gpt_unavailable")

        prompt = (
            "Extract invoice fields from OCR text and return STRICT JSON only. "
            "Return exactly these top-level keys: "
            "invoice_number, invoice_date, due_date, po_number, payment_terms, "
            "vendor_name, vendor_tax_id, customer_name, customer_tax_id, "
            "subtotal, tax, shipping, discounts, total, currency, order_items. "
            "For missing scalar fields use null. For missing order_items use []. "
            "order_items must be a list of objects with keys: "
            "line_no, description, qty, unit, unit_price, net_amount, tax_rate, gross_amount. "
            "Do not include markdown or extra keys."
        )

        try:
            response = self.openai_client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "You are a precise invoice field extraction engine."},
                    {"role": "user", "content": f"{prompt}\n\nOCR_TEXT:\n{ocr_text}"},
                ],
                temperature=0,
                max_tokens=1000,
            )

            content = (response.choices[0].message.content or "").strip()
            if self.is_policy_block(content):
                return self._extraction_error_fields("gpt_policy_block")

            return self._parse_json_or_error(content, "gpt_parse_error")

        except Exception as exc:
            if self.is_policy_block(str(exc)):
                return self._extraction_error_fields("gpt_policy_block")
            return self._extraction_error_fields("gpt_error", str(exc))

    def _ensure_qwen_loaded(self) -> bool:
        if self._qwen_model is not None and self._qwen_processor is not None:
            logger.info("Qwen already loaded model=%s device=%s", self.qwen_model_name, self._qwen_device)
            return True
        if self._qwen_load_error is not None:
            logger.error("Skipping Qwen load because previous load failed error=%s", self._qwen_load_error)
            return False
        if not self.qwen_model_dir.exists():
            self._qwen_load_error = f"Qwen model directory not found: {self.qwen_model_dir}"
            logger.error("Qwen model directory missing path=%s", self.qwen_model_dir)
            return False

        try:
            import torch
            from transformers import AutoModelForCausalLM, AutoModelForImageTextToText, AutoProcessor

            logger.info("Loading Qwen processor model_dir=%s", self.qwen_model_dir)
            self._qwen_processor = AutoProcessor.from_pretrained(
                str(self.qwen_model_dir),
                local_files_only=True,
                trust_remote_code=True,
            )
            model_kwargs = {
                "local_files_only": True,
                "trust_remote_code": True,
                "torch_dtype": "auto",
            }
            try:
                logger.info("Loading Qwen with AutoModelForImageTextToText model_dir=%s", self.qwen_model_dir)
                self._qwen_model = AutoModelForImageTextToText.from_pretrained(
                    str(self.qwen_model_dir),
                    **model_kwargs,
                )
            except Exception as image_text_exc:
                logger.exception(
                    "AutoModelForImageTextToText load failed; falling back to AutoModelForCausalLM model_dir=%s",
                    self.qwen_model_dir,
                )
                self._qwen_model = AutoModelForCausalLM.from_pretrained(
                    str(self.qwen_model_dir),
                    **model_kwargs,
                )
            device_errors = []
            for device in self._qwen_device_candidates():
                try:
                    logger.info("Moving Qwen model to device=%s", device)
                    self._qwen_model.to(device)
                    self._qwen_model.eval()
                    self._qwen_device = device
                    logger.info("Qwen loaded successfully model=%s device=%s", self.qwen_model_name, device)
                    return True
                except Exception as device_exc:
                    logger.exception("Qwen device activation failed device=%s", device)
                    device_errors.append(f"{device}: {device_exc}")
            self._qwen_load_error = "; ".join(device_errors) or "No usable Qwen device found."
            logger.error("Qwen load failed on all devices error=%s", self._qwen_load_error)
            return False
        except Exception as exc:
            self._qwen_load_error = str(exc)
            logger.exception("Qwen load failed model_dir=%s", self.qwen_model_dir)
            return False

    def extract_fields_qwen(self, ocr_text: str) -> Dict[str, Any]:
        logger.info(
            "Qwen extraction requested text_chars=%s model=%s model_dir=%s",
            len(ocr_text or ""),
            self.qwen_model_name,
            self.qwen_model_dir,
        )
        if not self._ensure_qwen_loaded():
            return self._extraction_error_fields("qwen_unavailable", self._qwen_load_error)

        prompt = (
            "Extract invoice fields from OCR text and return compact STRICT JSON only. "
            "Return exactly these top-level keys: "
            "invoice_number, invoice_date, due_date, po_number, payment_terms, "
            "vendor_name, vendor_tax_id, customer_name, customer_tax_id, "
            "subtotal, tax, shipping, discounts, total, currency, order_items. "
            "For missing scalar fields use null. For missing order_items use []. "
            "order_items must be a list of objects with keys: "
            "line_no, description, qty, unit, unit_price, net_amount, tax_rate, gross_amount. "
            "Do not include markdown, explanations, or extra keys."
        )

        try:
            import torch
            qwen_ocr_text = self._qwen_ocr_context(ocr_text)
            logger.info(
                "Qwen prompt prepared original_chars=%s prompt_ocr_chars=%s device=%s",
                len(ocr_text or ""),
                len(qwen_ocr_text),
                self._qwen_device,
            )

            messages = [
                {"role": "system", "content": "You are a precise invoice field extraction engine."},
                {"role": "user", "content": f"{prompt}\n\nOCR_TEXT:\n{qwen_ocr_text}"},
            ]
            chat_text = self._qwen_processor.apply_chat_template(
                messages,
                tokenize=False,
                add_generation_prompt=True,
            )
            inputs = self._qwen_processor(text=[chat_text], padding=True, return_tensors="pt")
            inputs = {
                k: (v.to(self._qwen_device) if hasattr(v, "to") else v)
                for k, v in inputs.items()
            }

            attempts = [
                ("initial", self.qwen_max_time_seconds, self.qwen_max_new_tokens),
                (
                    "partial_json_retry",
                    max(self.qwen_retry_max_time_seconds, self.qwen_max_time_seconds),
                    max(self.qwen_max_new_tokens, 768),
                ),
            ]
            last_content = ""
            for attempt_name, max_time, max_new_tokens in attempts:
                with torch.inference_mode():
                    generation_kwargs = {
                        "max_new_tokens": max_new_tokens,
                        "do_sample": False,
                        "num_beams": 1,
                        "use_cache": True,
                    }
                    if max_time > 0:
                        generation_kwargs["max_time"] = max_time
                    logger.info(
                        "Qwen generation attempt=%s max_time=%s max_new_tokens=%s",
                        attempt_name,
                        max_time,
                        max_new_tokens,
                    )
                    generated_ids = self._qwen_model.generate(
                        **inputs,
                        **generation_kwargs,
                    )

                generated_ids_trimmed = [
                    out_ids[len(in_ids) :]
                    for in_ids, out_ids in zip(inputs["input_ids"], generated_ids)
                ]
                content = self._qwen_processor.batch_decode(
                    generated_ids_trimmed,
                    skip_special_tokens=True,
                    clean_up_tokenization_spaces=False,
                )[0].strip()
                last_content = content
                logger.info(
                    "Qwen generation completed attempt=%s output_chars=%s output_preview=%r",
                    attempt_name,
                    len(content),
                    content[:500],
                )
                parsed = self._parse_model_json(content)
                if parsed is not None:
                    return self._normalize_extracted_fields(parsed)
                if attempt_name == "initial" and self._looks_like_partial_json(content):
                    logger.warning("Qwen output looks incomplete; retrying with longer generation window.")
                    continue
                break
            return self._extraction_error_fields("qwen_parse_error", last_content[:500])
        except Exception as exc:
            logger.exception("Qwen extraction failed")
            return self._extraction_error_fields("qwen_error", str(exc))

    def extract_fields_with_mode(self, ocr_text: str) -> Tuple[Dict[str, Any], str]:
        mode = self.field_extractor_mode
        logger.info("Field extraction dispatch mode=%s text_chars=%s", mode, len(ocr_text or ""))

        if mode == "gpt":
            fields = self.extract_fields_gpt4omini(ocr_text)
            logger.info("Field extraction completed provider=gpt error=%s", fields.get("extraction_error"))
            return fields, "gpt-4o-mini"

        if mode == "qwen":
            fields = self.extract_fields_qwen(ocr_text)
            logger.info("Field extraction completed provider=qwen error=%s", fields.get("extraction_error"))
            return fields, self.qwen_model_name

        # auto: GPT first, then Qwen local. No regex fallback.
        if self.openai_client is not None:
            gpt_fields = self.extract_fields_gpt4omini(ocr_text)
            if not gpt_fields.get("extraction_error"):
                logger.info("Field extraction auto selected provider=gpt")
                return gpt_fields, "gpt-4o-mini"
            logger.warning(
                "Field extraction auto GPT failed error=%s detail=%s; trying Qwen",
                gpt_fields.get("extraction_error"),
                gpt_fields.get("extraction_error_detail"),
            )

        qwen_fields = self.extract_fields_qwen(ocr_text)
        if not qwen_fields.get("extraction_error"):
            logger.info("Field extraction auto selected provider=qwen")
            return qwen_fields, self.qwen_model_name
        logger.error(
            "Field extraction failed provider=qwen error=%s detail=%s",
            qwen_fields.get("extraction_error"),
            qwen_fields.get("extraction_error_detail"),
        )
        return qwen_fields, self.qwen_model_name

    def ocr_jpg_upload(self, uploaded_file) -> Dict[str, Any]:
        logger.info("OCR JPG upload start filename=%s", getattr(uploaded_file, "name", None))
        file_bytes = np.asarray(bytearray(uploaded_file.read()), dtype=np.uint8)
        bgr = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        if bgr is None:
            logger.error("OCR JPG decode failed filename=%s bytes=%s", getattr(uploaded_file, "name", None), len(file_bytes))
            return {"status": "error", "error": "Image decode failed"}

        rgb = cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB)
        metrics = self.quality_metrics(rgb)
        text, conf, det, layout_meta = self.ocr_image_layout_aware(rgb)
        fields, mode = self.extract_fields_with_mode(text)
        logger.info(
            "OCR JPG upload complete filename=%s text_chars=%s confidence=%.4f detections=%s extraction_mode=%s extraction_error=%s",
            getattr(uploaded_file, "name", None),
            len(text or ""),
            conf,
            det,
            mode,
            fields.get("extraction_error"),
        )

        return {
            "status": "success",
            "type": "jpg",
            "avg_confidence": conf,
            "detections": det,
            "quality": metrics,
            "text": text,
            "fields": fields,
            "extraction_mode": mode,
            "layout": layout_meta,
        }

    def ocr_pdf_upload(self, uploaded_file) -> Dict[str, Any]:
        logger.info("OCR PDF upload start filename=%s", getattr(uploaded_file, "name", None))
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = Path(tmp.name)
        try:
            doc = pdfium.PdfDocument(str(tmp_path))
            logger.info("OCR PDF opened filename=%s pages=%s", getattr(uploaded_file, "name", None), len(doc))
            page_outputs: List[Dict[str, Any]] = []
            all_text: List[str] = []
            all_conf: List[float] = []

            for i in range(len(doc)):
                arr = np.array(doc[i].render(scale=2.0).to_pil().convert("RGB"))
                text, conf, det, layout_meta = self.ocr_image_layout_aware(arr)
                page_outputs.append(
                    {
                        "page": i + 1,
                        "avg_confidence": conf,
                        "detections": det,
                        "layout_status": layout_meta.get("layout_status"),
                        "layout_regions": layout_meta.get("layout_regions", 0),
                        "layout_crop_chunks": layout_meta.get("layout_crop_chunks", 0),
                    }
                )
                all_text.append(f"=== PAGE {i + 1} ===\n{text}")
                if det > 0:
                    all_conf.append(conf)

            merged = "\n\n".join(all_text)
            fields, mode = self.extract_fields_with_mode(merged)
            logger.info(
                "OCR PDF upload complete filename=%s pages=%s text_chars=%s extraction_mode=%s extraction_error=%s",
                getattr(uploaded_file, "name", None),
                len(doc),
                len(merged or ""),
                mode,
                fields.get("extraction_error"),
            )

            return {
                "status": "success",
                "type": "pdf",
                "pages": len(doc),
                "avg_confidence": float(sum(all_conf) / len(all_conf)) if all_conf else 0.0,
                "page_stats": page_outputs,
                "layout_summary": {
                    "regions_total": int(sum(p.get("layout_regions", 0) for p in page_outputs)),
                    "crop_chunks_total": int(sum(p.get("layout_crop_chunks", 0) for p in page_outputs)),
                },
                "text": merged,
                "fields": fields,
                "extraction_mode": mode,
            }
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass

    def ocr_docx_upload(self, uploaded_file) -> Dict[str, Any]:
        """Process a .docx file: extract native text + OCR any embedded images."""
        import io
        import zipfile
        import docx

        logger.info("OCR DOCX upload start filename=%s", getattr(uploaded_file, "name", None))
        raw = uploaded_file.read()

        # --- 1. Native text via python-docx (paragraphs + table cells) ---
        doc_obj = docx.Document(io.BytesIO(raw))
        native_parts: List[str] = []
        for para in doc_obj.paragraphs:
            t = para.text.strip()
            if t:
                native_parts.append(t)
        for table in doc_obj.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    native_parts.append(" | ".join(row_cells))
        native_text = "\n".join(native_parts)

        # --- 2. Embedded images via ZIP (word/media/*) → EasyOCR ---
        image_texts: List[str] = []
        image_stats: List[Dict[str, Any]] = []
        supported_img_exts = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"}

        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            media_files = [
                n for n in z.namelist()
                if n.startswith("word/media/")
                and Path(n).suffix.lower() in supported_img_exts
            ]
            for idx, media_name in enumerate(media_files):
                img_bytes = z.read(media_name)
                arr = cv2.imdecode(np.frombuffer(img_bytes, np.uint8), cv2.IMREAD_COLOR)
                if arr is None:
                    continue
                arr_rgb = cv2.cvtColor(arr, cv2.COLOR_BGR2RGB)
                text, conf, det, layout_meta = self.ocr_image_layout_aware(arr_rgb)
                image_texts.append(f"=== EMBEDDED IMAGE {idx + 1} ({Path(media_name).name}) ===\n{text}")
                image_stats.append({
                    "image": Path(media_name).name,
                    "avg_confidence": conf,
                    "detections": det,
                    "layout_status": layout_meta.get("layout_status"),
                    "layout_regions": layout_meta.get("layout_regions", 0),
                    "layout_crop_chunks": layout_meta.get("layout_crop_chunks", 0),
                })

        # --- 3. Combine and extract fields ---
        sections: List[str] = []
        if native_text.strip():
            sections.append(f"=== DOCUMENT TEXT ===\n{native_text}")
        sections.extend(image_texts)
        merged = "\n\n".join(sections) if sections else native_text

        fields, mode = self.extract_fields_with_mode(merged)
        logger.info(
            "OCR DOCX upload complete filename=%s native_lines=%s embedded_images=%s text_chars=%s extraction_mode=%s extraction_error=%s",
            getattr(uploaded_file, "name", None),
            len(native_parts),
            len(media_files),
            len(merged or ""),
            mode,
            fields.get("extraction_error"),
        )

        return {
            "status": "success",
            "type": "docx",
            "native_text_lines": len(native_parts),
            "embedded_images": len(media_files),
            "image_stats": image_stats,
            "text": merged,
            "fields": fields,
            "extraction_mode": mode,
        }

    def process_upload(self, uploaded_file) -> Dict[str, Any]:
        suffix = Path(uploaded_file.name).suffix.lower()
        logger.info("OCR process_upload dispatch filename=%s suffix=%s", getattr(uploaded_file, "name", None), suffix)
        try:
            if suffix == ".pdf":
                return self.ocr_pdf_upload(uploaded_file)
            if suffix == ".docx":
                return self.ocr_docx_upload(uploaded_file)
            return self.ocr_jpg_upload(uploaded_file)
        except Exception:
            logger.exception("OCR process_upload failed filename=%s suffix=%s", getattr(uploaded_file, "name", None), suffix)
            raise
